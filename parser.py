import win32com.client as win32
from docx import Document
import os


def convert_doc_to_docx(doc_path):
    if not os.path.exists(doc_path):
        print(f"Error: File not found at {doc_path}")
        return None

    base_dir = os.path.dirname(os.path.abspath(doc_path))
    file_name = os.path.splitext(os.path.basename(doc_path))[0]
    docx_path = os.path.join(base_dir, f"{file_name}.docx")

    word = None
    try:
        word = win32.Dispatch("Word.Application")
        word.visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(docx_path, FileFormat=16)
        doc.Close()
        print(f"Successfully converted '{doc_path}' to '{docx_path}'")
        return docx_path
    except Exception as e:
        print(f"Error during conversion: {e}")
        return None
    finally:
        if word:
            word.Quit()


def get_cell_value(table, row, col):
    try:
        cell = table.cell(row, col)
        return cell.text.strip()
    except IndexError:
        print(f"Error: Row {row} or column {col} is out of bounds for the table.")
        return None


def is_merged(table, row, col):
    cell = table.cell(row, col)
    if row > 0 and cell._tc == table.cell(row - 1, col)._tc:
        return True
    if col > 0 and cell._tc == table.cell(row, col - 1)._tc:
        return True
    return False


def extract_days(table):
    days = []
    start_row = 1
    for r_idx in range(start_row, len(table.rows)):
        day_text = get_cell_value(table, r_idx, 0)
        if day_text and day_text not in days:
            days.append(day_text)
    return days


def extract_groups(table):
    groups = []
    header_row_index = 0
    start_col = 2
    if len(table.rows) > header_row_index:
        header_row = table.rows[header_row_index]
        for c_idx in range(start_col, len(header_row.cells)):
            group_name = get_cell_value(table, header_row_index, c_idx)
            if group_name:
                groups.append(group_name)
    return groups


def extract_class_times(table):
    class_times = []
    start_row = 1
    col_index = 1
    for r_idx in range(start_row, len(table.rows)):
        class_time_text = get_cell_value(table, r_idx, col_index)
        if class_time_text and class_time_text not in class_times:
            class_times.append(class_time_text)
    return class_times


def extract_schedule_by_day(table):
    schedule = {}
    current_day = None
    start_row = 1
    day_col, class_time_col = 0, 1

    for r_idx in range(start_row, len(table.rows)):
        day_text = get_cell_value(table, r_idx, day_col)
        if day_text:
            current_day = day_text
            if current_day not in schedule:
                schedule[current_day] = []

        if current_day:
            class_time_text = get_cell_value(table, r_idx, class_time_col)
            if class_time_text and class_time_text not in schedule[current_day]:
                schedule[current_day].append(class_time_text)
    return schedule


def extract_schedule_for_all_groups(table):
    groups = extract_groups(table)
    group_schedules = {group: {} for group in groups}

    processed_cells = set()
    processed_denominator_rows = set()

    start_row = 1
    day_col, class_time_col = 0, 1

    for r_idx in range(start_row, len(table.rows)):
        day_text = get_cell_value(table, r_idx, day_col)
        current_day = day_text if day_text else current_day

        if not current_day:
            continue

        for group in groups:
            if current_day not in group_schedules[group]:
                group_schedules[group][current_day] = []

        class_time = get_cell_value(table, r_idx, class_time_col)
        if not class_time:
            continue

        for i, group in enumerate(groups):
            group_col_idx = i + 2
            if (r_idx, group_col_idx) in processed_cells:
                continue

            class_details_num = get_cell_value(table, r_idx, group_col_idx)
            class_details_den = ""

            is_denominator_week_present = (r_idx + 1 < len(table.rows) and
                                           not get_cell_value(table, r_idx + 1, day_col) and
                                           not get_cell_value(table, r_idx + 1, class_time_col))

            if is_denominator_week_present and r_idx not in processed_denominator_rows:
                class_details_den = get_cell_value(table, r_idx + 1, group_col_idx)
                processed_denominator_rows.add(r_idx + 1)

            if not class_details_num and not class_details_den:
                processed_cells.add((r_idx, group_col_idx))
                continue

            affected_groups = [group]
            affected_cols = [group_col_idx]

            next_col_idx = group_col_idx + 1
            while next_col_idx < len(table.columns) and is_merged(table, r_idx, next_col_idx):
                if (r_idx, next_col_idx) not in processed_cells:
                    affected_groups.append(groups[next_col_idx - 2])
                    affected_cols.append(next_col_idx)
                next_col_idx += 1

            if class_details_num and class_details_num == class_details_den:
                final_details = class_details_num
            elif class_details_num and class_details_den:
                final_details = f"Numerator: {class_details_num}, Denominator: {class_details_den}"
            elif class_details_num:
                final_details = f"Numerator: {class_details_num}"
            else:
                final_details = f"Denominator: {class_details_den}"

            schedule_entry = {"time": class_time, "details": final_details}

            for grp in affected_groups:
                if schedule_entry not in group_schedules[grp][current_day]:
                    group_schedules[grp][current_day].append(schedule_entry)

            for col in affected_cols:
                processed_cells.add((r_idx, col))
                if is_denominator_week_present:
                    processed_cells.add((r_idx + 1, col))

    return group_schedules


def has_complex_columns(table, row_idx, start_col_idx, end_col_idx):
    complex_cols = []

    for col in range(start_col_idx, end_col_idx + 1):
        if col < len(table.columns):
            cell = table.cell(row_idx, col)
            try:
                grid_span = cell._tc.tcPr.gridSpan.val
                if grid_span > 1:
                    complex_cols.append(col)
            except AttributeError:
                pass

            for below_row in range(row_idx + 1, len(table.rows)):
                if cell._tc == table.cell(below_row, col)._tc:
                    complex_cols.append(col)
                    break

    return complex_cols


def has_complex_rows(table, col_idx, start_row_idx, end_row_idx):
    complex_rows = []

    for row in range(start_row_idx, end_row_idx + 1):
        if row < len(table.rows) and col_idx < len(table.columns):
            cell = table.cell(row, col_idx)

            for below_row in range(row + 1, len(table.rows)):
                if cell._tc == table.cell(below_row, col_idx)._tc:
                    complex_rows.append(row)
                    break

    return complex_rows


if __name__ == "__main__":
    # doc_path = "D://schedule-parser//schedules//Розклад_2022_2023_І_семестр_Інф_І_курс.doc"  # Replace with your .doc file path
    # docx_path = convert_doc_to_docx(doc_path)
    # if docx_path:
    #     print(f"Converted file is located at: {docx_path}")
    #     document = Document(docx_path)
    #     if document.tables:
    #         first_table = document.tables[0]
    #         cell_value = get_cell_value(first_table, 4, 2)
    #         if cell_value is not None:
    #             print(f"Value from table 0, row 5, col 2: '{cell_value}'")
    #     else:
    #         print("No tables found in the document.")
    # else:
    #     print("Conversion failed.")
    docx_path1 = "D://schedule-parser//schedules//Розклад_2022_2023_І_семестр_Інф_І_курс.docx"
    document = Document(docx_path1)
    first_table = document.tables[0]
    # cell_value = get_cell_value(first_table, 1, 6)
    # print(cell_value)

    row_to_check = 0
    col_to_check = 1

    cell_value = get_cell_value(first_table, row_to_check, col_to_check)
    print(f"Value at ({row_to_check}, {col_to_check}): '{cell_value}'")

    if document.tables:
        first_table = document.tables[0]

        all_schedules = extract_schedule_for_all_groups(first_table)

        target_group = "ПМІ – 11с"
        group_schedule = all_schedules.get(target_group)

        if group_schedule:
            print(f"Schedule for Group: {target_group}")
            for day, classes in group_schedule.items():
                print(f"\n--- {day} ---")
                if classes:
                    for c in classes:
                        print(f"  {c['time']}: {c['details']}")
                else:
                    print("  No classes scheduled.")
            print("-" * 20)
        else:
            print(f"Could not retrieve schedule for group '{target_group}'.")

    else:
        print("No tables found in the document.")
